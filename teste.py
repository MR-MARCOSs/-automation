import pandas as pd
from docx import Document

# Função para preencher o contrato com os dados do cliente
def preencher_contrato(nome, endereco, cpf, valor, servico, data_contrato):
    # Carregar o template do contrato
    doc = Document()

    # Conteúdo do contrato
    doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS', level=1)
    
    doc.add_paragraph(f"Contratante: {nome}")
    doc.add_paragraph(f"Endereço: {endereco}")
    doc.add_paragraph(f"CPF/CNPJ: {cpf}")

    # Dados do contratado (você pode modificar conforme necessário)
    doc.add_paragraph("\nContratado: Nome do Contratado")
    doc.add_paragraph("Endereço: Endereço do Contratado")
    doc.add_paragraph("CPF/CNPJ: Documento do Contratado")

    # Cláusulas do contrato
    doc.add_heading('CLÁUSULAS DO CONTRATO', level=2)
    doc.add_paragraph(f"1. OBJETO: O presente contrato tem por objeto a prestação dos seguintes serviços: {servico}.")
    doc.add_paragraph("2. OBRIGAÇÕES DO CONTRATANTE: [Descrever as obrigações do contratante].")
    doc.add_paragraph("3. OBRIGAÇÕES DO CONTRATADO: [Descrever as obrigações do contratado].")
    doc.add_paragraph(f"4. VALORES E CONDIÇÕES DE PAGAMENTO: O valor total dos serviços prestados será de R$ {valor}, pagos conforme [Condições de pagamento].")
    doc.add_paragraph(f"5. PRAZO: O contrato terá duração de [Duração do contrato], com início em {data_contrato} e término em [Data de término].")
    doc.add_paragraph("6. RESCISÃO: O contrato poderá ser rescindido por qualquer das partes mediante aviso prévio de [Prazo] dias.")
    doc.add_paragraph("7. DISPOSIÇÕES FINAIS: [Outras disposições importantes].")

    # Assinaturas
    doc.add_paragraph("\n\n_______________________________")
    doc.add_paragraph("Assinatura do Contratante")

    doc.add_paragraph("\n_______________________________")
    doc.add_paragraph("Assinatura do Contratado")

    # Salvar o documento com o nome do cliente
    doc.save(f"Contrato_{nome.replace(' ', '_')}.docx")

# Ler os dados do Excel
df = pd.read_excel(r'C:\Users\marco\Desktop\codes\automatizacao\Clientes_Contrato.xlsx')

# Loop pelos dados e preencher contratos
for index, row in df.iterrows():
    preencher_contrato(row['Nome do Cliente'], row['Endereço'], row['CPF/CNPJ'], 
                       row['Valor do Contrato (R$)'], row['Serviços Contratados'], 
                       row['Data do Contrato'])
